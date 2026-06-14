"""docx source tool: ContentBlocks per heading-grouped section and table row."""
from __future__ import annotations

from pathlib import Path

from app.core.ckm import ContentBlock

from ._base import OnBlock, _emit, _slug


def extract_docx(path: Path, on_block: OnBlock = None) -> list[ContentBlock]:
    from docx import Document

    doc = Document(str(path))
    blocks: list[ContentBlock] = []

    # --- paragraphs grouped under headings ---
    current_title = path.stem
    buffer: list[str] = []
    section = 0

    def flush() -> None:
        nonlocal buffer, section
        text = "\n".join(buffer).strip()
        if text:
            _emit(blocks, ContentBlock(
                id=_slug(path.stem, "s", section),
                source=path.name,
                modality="text",
                title=current_title[:80],
                text=text,
                metadata={"section": section},
            ), on_block)
            section += 1
        buffer = []

    for para in doc.paragraphs:
        txt = para.text.strip()
        if not txt:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading") or style == "title":
            flush()
            current_title = txt
            _emit(blocks, ContentBlock(
                id=_slug(path.stem, "h", section),
                source=path.name,
                modality="heading",
                title=txt[:80],
                text=txt,
                metadata={"section": section},
            ), on_block)
        else:
            buffer.append(txt)
    flush()

    # --- tables: one block per row ---
    for ti, table in enumerate(doc.tables):
        rows = table.rows
        if not rows:
            continue
        header = [c.text.strip() for c in rows[0].cells]
        for ri, row in enumerate(rows[1:], start=1):
            cells = [c.text.strip() for c in row.cells]
            if not any(cells):
                continue
            pairs = {header[j]: cells[j] for j in range(min(len(header), len(cells)))
                     if header[j] and cells[j]}
            _emit(blocks, ContentBlock(
                id=_slug(path.stem, "t", ti, ri),
                source=path.name,
                modality="table_row",
                title=f"Table {ti + 1} row {ri}",
                text=" | ".join(f"{k}: {v}" for k, v in pairs.items()) or " | ".join(cells),
                metadata={"table": ti, "fields": pairs},
            ), on_block)
    return blocks